import markdown
from docx import Document
from bs4 import BeautifulSoup

def markdown_to_word(markdown_file, word_file):
    # Reading the Markdown file
    with open(markdown_file, 'r', encoding='utf-8') as file:
        markdown_content = file.read()

    # Converting Markdown to HTML
    html_content = markdown.markdown(markdown_content)

    # Creating a new Word Document
    doc = Document()

    # Converting HTML to text and add it to the Word Document
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Adding content to the Word Document
    for element in soup:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'p':
            paragraph = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong':
                    paragraph.add_run(child.text).bold = True
                elif child.name == 'em':
                    paragraph.add_run(child.text).italic = True
                else:
                    paragraph.add_run(child)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
    
    doc.save(word_file)

markdown_file = input("Enter the path to the Markdown file (e.g., README.md): ")
word_file = input("Enter the path for the output Word file (e.g., README.docx): ")

markdown_to_word(markdown_file, word_file)
